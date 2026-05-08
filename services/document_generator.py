"""
services/document_generator.py
================================
Document generation service for the career-agent-email-cover project.

Produces two output formats per job application:
  * Markdown (.md)  — human-readable, easy to paste into email clients.
  * DOCX (.docx)    — professional format for formal applications.

Uses ``python-docx`` for DOCX generation.  No external LibreOffice or
headless browser is required, so this works cleanly inside GitHub Actions.

Usage
-----
    from services.document_generator import save_markdown, save_docx, build_output_paths
    from pathlib import Path

    paths = build_output_paths(Path("output"), "Acme Corp", "JOB-123")
    save_markdown(email_text, paths["email_md"])
    save_docx(cover_letter_text, paths["cover_letter_docx"], title="Cover Letter")
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from services.logger import setup_logger

logger = setup_logger(__name__)


# ======================================================================== #
# Public API                                                                 #
# ======================================================================== #

def save_markdown(content: str, file_path: Path) -> None:
    """
    Write ``content`` to a UTF-8 encoded Markdown file.

    Parameters
    ----------
    content:
        Full Markdown string (may include ``#`` headings, ``**bold**``, etc.).
    file_path:
        Destination path; parent directory must exist.
    """
    file_path.write_text(content, encoding="utf-8")
    logger.info(f"Saved Markdown: {file_path}")


def save_docx(content: str, file_path: Path, title: str = "") -> None:
    """
    Write ``content`` to a DOCX file with clean professional formatting.

    Each newline in ``content`` becomes a separate paragraph.  An optional
    centred heading is added at the top when ``title`` is provided.

    Parameters
    ----------
    content:
        Plain text content of the document (no Markdown syntax).
    file_path:
        Destination path; parent directory must exist.
    title:
        Optional document title rendered as a centred Heading 1.
    """
    doc = Document()

    # --- Page margins: 1 inch top/bottom, 1.25 inch left/right (professional default) ---
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # --- Optional centred title ---
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add a blank line after the title for visual breathing room
        doc.add_paragraph()

    # --- Body: one paragraph per line ---
    for line in content.split("\n"):
        para = doc.add_paragraph(line)
        # Apply a clean readable font to the default "Normal" style on each para
        run = para.runs[0] if para.runs else para.add_run()
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    doc.save(str(file_path))
    logger.info(f"Saved DOCX: {file_path}")


def build_output_paths(output_dir: Path, company: str, job_id: str) -> dict[str, Path]:
    """
    Compute and create all output file paths for one job application.

    Files are organised under ``output_dir / <safe_company_name>/``.

    Parameters
    ----------
    output_dir:
        Root output directory (from config.output_dir).
    company:
        Company name — sanitised for use in directory and file names.
    job_id:
        Job posting ID — appended to file names to avoid collisions.

    Returns
    -------
    dict[str, Path]
        Keys: ``email_md``, ``cover_letter_md``, ``cover_letter_docx``.
    """
    # Sanitise: keep alphanumerics, spaces, hyphens, underscores; replace rest
    safe_company = re.sub(r"[^\w\s\-]", "_", company).strip()
    safe_company = re.sub(r"\s+", "_", safe_company)  # spaces → underscores

    # Use job_id in the filename prefix for uniqueness; fall back to company only
    prefix = f"{safe_company}_{job_id}" if job_id else safe_company

    # Create per-company sub-directory
    job_dir = output_dir / safe_company
    job_dir.mkdir(parents=True, exist_ok=True)

    return {
        "email_md": job_dir / f"{prefix}_recruiter_email.md",
        "cover_letter_md": job_dir / f"{prefix}_cover_letter.md",
        "cover_letter_docx": job_dir / f"{prefix}_cover_letter.docx",
    }
